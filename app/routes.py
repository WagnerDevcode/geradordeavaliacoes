import os
import csv
import logging
from docx import Document
from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    send_file,
    url_for,
    flash,
    jsonify,
)
from fpdf import FPDF
from docx.shared import Inches, Pt
from werkzeug.utils import secure_filename
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


from docx.oxml.ns import qn

# Configuração do logger
logging.basicConfig(level=logging.DEBUG)

main = Blueprint("main", __name__)

# Diretórios e caminhos
AVALIACOES_DIR = "avaliacoes"  # Diretório para armazenar as avaliações geradas

ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "csv"}
CSV_FILE_PATH = "questoes.csv"
IMAGES_DIR = os.path.join(os.getcwd(), "app", "static", "header_images")
AVALIACOES_DIR = os.path.join(
    os.getcwd(), "avaliacoes"
)  # Caminho relativo à pasta atual


# Diretórios e caminhos
AVALIACOES_DIR = os.path.join(os.getcwd(), "avaliacoes")
PDF_FOLDER = os.path.join(os.getcwd(), "app" "static", "avaliacoes", "pdf")
DOCX_FOLDER = os.path.join(os.getcwd(), "app" "static", "avaliacoes", "docx")
IMAGES_DIR = os.path.join(os.getcwd(), "app" "static", "header_images")
UPLOAD_FOLDER = os.path.join(os.getcwd(), "app", "static", "uploads")

os.makedirs(PDF_FOLDER, exist_ok=True)
os.makedirs(DOCX_FOLDER, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

HEADER_PADRAO = [
    "ID",
    "Código BNCC",
    "Disciplina",
    "Assunto",
    "Autor",
    "Enunciado",
    "Alternativa A",
    "Alternativa B",
    "Alternativa C",
    "Alternativa D",
    "Alternativa E",
    "Gabarito",
    "Pontuação",
    "Imagem",
]


def allowed_file(filename):
    """Verifica se o arquivo tem uma extensão permitida."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def verificar_cabecalho_csv():
    """Verifica se o cabeçalho do CSV está correto e o corrige se necessário."""
    if os.path.exists(CSV_FILE_PATH):
        with open(CSV_FILE_PATH, mode="r", encoding="utf-8") as csvfile:
            reader = csv.reader(csvfile)
            cabecalho_existente = next(reader, None)

            if cabecalho_existente != HEADER_PADRAO:
                corrigir_cabecalho_csv()


def corrigir_cabecalho_csv():
    """Corrige o cabeçalho do CSV caso esteja ausente ou incorreto."""
    with open(CSV_FILE_PATH, mode="r", encoding="utf-8") as csvfile:
        linhas = csvfile.readlines()

    with open(CSV_FILE_PATH, mode="w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(HEADER_PADRAO)

        if len(linhas) > 1:
            csvfile.writelines(linhas[1:])


def initialize_csv():
    """Inicializa o arquivo CSV com o cabeçalho correto caso ele não exista."""
    if not os.path.exists(CSV_FILE_PATH):
        with open(CSV_FILE_PATH, mode="w", newline="", encoding="utf-8") as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(HEADER_PADRAO)


def get_next_id():
    """Obtém o próximo ID único para uma nova questão."""
    if not os.path.exists(CSV_FILE_PATH):
        return 1
    with open(CSV_FILE_PATH, mode="r", encoding="utf-8") as csvfile:
        reader = csv.reader(csvfile)
        next(reader)
        ids = [int(row[0]) for row in reader if row[0].isdigit()]
        return max(ids) + 1 if ids else 1


def ler_questoes():
    """Lê as questões do arquivo CSV com codificação correta."""
    questoes = []
    try:
        with open(CSV_FILE_PATH, mode="r", encoding="utf-8") as file:
            reader = csv.DictReader(file)
            questoes = [row for row in reader]
    except Exception as e:
        print(f"Erro ao ler o arquivo CSV: {e}")
    return questoes


def garantir_pasta_avaliacoes():
    """Garante que a pasta de avaliações exista."""
    os.makedirs(AVALIACOES_DIR, exist_ok=True)


def gerar_csv_avaliacao(ids_selecionados, nome_avaliacao):
    """Função para gerar o arquivo CSV de avaliação."""
    try:
        # Carrega todas as questões
        questoes = ler_questoes()

        # Filtra as questões selecionadas
        questoes_selecionadas = [q for q in questoes if q["ID"] in ids_selecionados]

        # Verifica se há questões selecionadas
        if not questoes_selecionadas:
            print("Nenhuma questão foi selecionada.")
            return None

        # Gera o caminho completo para salvar o arquivo CSV
        caminho_arquivo = os.path.join(AVALIACOES_DIR, f"{nome_avaliacao}.csv")

        # Cria o arquivo CSV
        with open(caminho_arquivo, mode="w", newline="", encoding="utf-8") as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(
                [
                    "ID",
                    "Código BNCC",
                    "Disciplina",
                    "Assunto",
                    "Autor",
                    "Enunciado",
                    "Alternativa A",
                    "Alternativa B",
                    "Alternativa C",
                    "Alternativa D",
                    "Alternativa E",
                    "Pontuação",
                    "Imagem",
                ]
            )

            # Escreve cada questão selecionada no arquivo CSV
            for questao in questoes_selecionadas:
                writer.writerow(
                    [
                        questao["ID"],
                        questao["Código BNCC"],
                        questao["Disciplina"],
                        questao["Assunto"],
                        questao["Autor"],
                        questao["Enunciado"],
                        questao["Alternativa A"],
                        questao["Alternativa B"],
                        questao["Alternativa C"],
                        questao["Alternativa D"],
                        questao["Alternativa E"],
                        questao["Pontuação"],
                        questao["Imagem"],
                    ]
                )

        # Retorna o nome do arquivo gerado
        return f"{nome_avaliacao}.csv"

    except Exception as e:
        print(f"Erro ao gerar o CSV: {e}")
        return None


def ler_questoes_csv(filepath):
    """Lê as questões do arquivo CSV."""
    with open(filepath, newline="") as csvfile:
        reader = csv.reader(csvfile)
        questoes = [row[0] for row in reader]  # Ajuste conforme a estrutura do CSV
    return questoes


@main.route("/")
def index():
    return render_template("index.html")


@main.route("/cadastro", methods=["GET", "POST"])
def cadastro():
    """Rota para cadastro de novas questões."""
    if request.method == "POST":
        # Captura os dados do formulário
        codigo_bncc = request.form["codigo_bncc"]
        disciplina = request.form["disciplina"]
        assunto = request.form["assunto"]
        autor = request.form["autor"]
        enunciado = request.form["enunciado"]
        alternativa_a = request.form["alternativa_a"]
        alternativa_b = request.form["alternativa_b"]
        alternativa_c = request.form["alternativa_c"]
        alternativa_d = request.form["alternativa_d"]
        alternativa_e = request.form["alternativa_e"]
        gabarito = request.form["gabarito"]
        pontuacao = request.form["pontuacao"]

        # Gera o próximo ID
        id_questao = get_next_id()

        # Processamento da imagem
        imagem = request.files.get("imagem")
        imagem_filename = None
        if imagem and allowed_file(imagem.filename):
            imagem_filename = secure_filename(imagem.filename)
            image_path = os.path.join(UPLOAD_FOLDER, imagem_filename)
            imagem.save(image_path)

        # Adiciona a nova questão ao CSV
        with open(CSV_FILE_PATH, mode="a", newline="", encoding="utf-8") as file:
            writer = csv.writer(file)
            writer.writerow(
                [
                    id_questao,
                    codigo_bncc,
                    disciplina,
                    assunto,
                    autor,
                    enunciado,
                    alternativa_a,
                    alternativa_b,
                    alternativa_c,
                    alternativa_d,
                    alternativa_e,
                    gabarito,
                    pontuacao,
                    imagem_filename,
                ]
            )

        flash("Questão cadastrada com sucesso!", "success")
        return redirect(url_for("main.index"))

    return render_template("cadastro.html")


@main.route("/visualizador", methods=["GET", "POST"])
def visualizador():
    """Rota para visualização e geração de avaliações."""
    if request.method == "POST":
        try:
            # Corrige para buscar os IDs com o nome correto
            ids_selecionados = request.form.getlist("questao_ids")
            nome_avaliacao = request.form.get("nome_avaliacao")

            print(f"IDs selecionados: {ids_selecionados}")
            print(f"Nome da avaliação: {nome_avaliacao}")

            # Gera o CSV usando os IDs e o nome da avaliação fornecidos
            arquivo_csv = gerar_csv_avaliacao(ids_selecionados, nome_avaliacao)
            if not arquivo_csv:
                return jsonify(
                    {"status": "error", "message": "Erro ao gerar a avaliação CSV."}
                )

            return jsonify({"status": "success", "arquivo": arquivo_csv})

        except Exception as e:
            print(f"Erro no servidor: {e}")
            return jsonify({"status": "error", "message": str(e)})

    questoes = ler_questoes()
    return render_template("visualizador.html", questoes=questoes)


# gerar pdf e docx


# Definição da rota para exibir e gerar avaliações
@main.route("/minhas_avaliacoes", methods=["GET", "POST"])
def minhas_avaliacoes():
    """Rota para exibir as avaliações e permitir gerar PDFs, DOCXs, e atualizar a imagem de cabeçalho."""
    if request.method == "POST":
        nome_avaliacao = request.form.get("nome_avaliacao")
        formato = request.form.get("formato")
        imagem = request.files.get("imagem")  # Receber a imagem do cabeçalho

        if not nome_avaliacao:
            return jsonify(
                {"status": "error", "message": "Nome da avaliação é obrigatório."}
            )

        arquivo_csv = os.path.join(AVALIACOES_DIR, f"{nome_avaliacao}.csv")
        if not os.path.exists(arquivo_csv):
            return jsonify({"status": "error", "message": "Avaliação não encontrada."})

        if imagem:
            # Salvar a imagem do cabeçalho associada ao nome da avaliação
            caminho_imagem = os.path.join(IMAGES_DIR, f"{nome_avaliacao}.png")
            imagem.save(caminho_imagem)
            return jsonify(
                {"status": "success", "message": "Imagem atualizada com sucesso."}
            )

        # Processar a geração do arquivo PDF ou DOCX
        if formato == "pdf":
            arquivo_pdf = gerar_pdf(arquivo_csv, nome_avaliacao)
            return send_file(arquivo_pdf, as_attachment=True)
        elif formato == "docx":
            arquivo_docx = gerar_docx(arquivo_csv, nome_avaliacao)
            return send_file(arquivo_docx, as_attachment=True)
        else:
            return jsonify({"status": "error", "message": "Formato inválido."})

    avaliacoes = ler_avaliacoes()
    return render_template("minhas_avaliacoes.html", avaliacoes=avaliacoes)


# Função para ler avaliações
def ler_avaliacoes():
    """Lê as avaliações do diretório e retorna uma lista com informações das avaliações."""
    avaliacoes = []
    try:
        for filename in os.listdir(AVALIACOES_DIR):
            if filename.endswith(".csv"):
                avaliacoes.append(
                    {
                        "nome": filename.replace(".csv", ""),
                        "imagem": os.path.join(
                            IMAGES_DIR, filename.replace(".csv", ".png")
                        ),  # Caminho da imagem associada
                    }
                )
    except Exception as e:
        print(f"Erro ao ler o diretório de avaliações: {e}")
    return avaliacoes


def gerar_pdf(arquivo_csv, nome_avaliacao):
    """Gera um arquivo PDF a partir do arquivo CSV de avaliação com formatação padrão ENEM."""
    pdf = FPDF()
    margem = 10  # Margem para a página
    largura_util = 210 - 2 * margem  # Largura utilizável da página A4 com margem
    pdf.set_auto_page_break(auto=True, margin=margem)

    def adicionar_cabecalho():
        """Adiciona o cabeçalho com imagem e título na primeira página."""
        cabecalho_img = os.path.join(IMAGES_DIR, f"{nome_avaliacao}.png")
        if os.path.exists(cabecalho_img):
            pdf.image(cabecalho_img, x=10, y=8, w=190)

        pdf.ln(65)  # Espaço menor abaixo da imagem de cabeçalho

        # Aumenta o tamanho do título
        pdf.set_font("Arial", style="B", size=14)
        pdf.cell(
            0,
            12,
            nome_avaliacao.encode("latin-1", "replace").decode("latin-1"),
            ln=True,
            align="C",
        )
        pdf.ln(0)  # Espaço ajustado após o título (menor)

    def adicionar_pagina(iniciar_apos_titulo=False):
        """Adiciona uma nova página e ajusta a margem superior."""
        pdf.add_page()
        if pdf.page_no() == 1:
            adicionar_cabecalho()
            # Define a posição inicial para ambas as colunas logo abaixo do título
            return pdf.get_y() + 2  # Espaço menor entre título e conteúdo
        else:
            # Nas outras páginas, começa no topo com a margem normal
            return margem

    def adicionar_borda_pagina():
        """Adiciona uma borda ao redor da página inteira."""
        pdf.set_draw_color(0, 0, 0)
        pdf.rect(x=margem / 2, y=margem / 2, w=210 - margem, h=297 - margem)

    def adicionar_linha_divisoria(y_inicial):
        """Adiciona uma linha divisória no meio da página, centralizada."""
        pdf.set_draw_color(0, 0, 0)
        x_start_right = margem + col_width + 5  # Centraliza a linha divisória
        pdf.line(x_start_right, y_inicial, x_start_right, 297 - margem)

    # Definições de layout
    col_width = (largura_util - 10) / 2  # Ajuste o cálculo para espaço entre colunas
    line_height = 7  # Altura de linha
    x_start_left = margem
    x_start_right = x_start_left + col_width + 10  # Espaço de 10mm entre as colunas
    current_x = x_start_left
    current_y = margem
    is_left_column = True

    # Adiciona a primeira página, o cabeçalho, e define o ponto inicial para o conteúdo
    current_y = adicionar_pagina(iniciar_apos_titulo=True)
    y_inicial_primeira_pagina = current_y  # Posição inicial para a segunda coluna
    adicionar_borda_pagina()
    adicionar_linha_divisoria(y_inicial_primeira_pagina)

    # Diretório de uploads de imagens
    upload_dir = os.path.join("static", "uploads")

    def ajustar_coluna_ou_pagina(altura_conteudo):
        """Ajusta para a próxima coluna ou página se necessário."""
        nonlocal current_x, current_y, is_left_column

        # Verifica se o conteúdo excede o limite da página
        if current_y + altura_conteudo > 297 - margem:
            # Se estiver na coluna esquerda, passa para a direita
            if is_left_column:
                current_x = x_start_right
                current_y = y_inicial_primeira_pagina if pdf.page_no() == 1 else margem
                is_left_column = False
            else:
                # Se a coluna direita está cheia, passa para a próxima página e reinicia na coluna esquerda
                current_y = adicionar_pagina(iniciar_apos_titulo=False)
                adicionar_borda_pagina()
                adicionar_linha_divisoria(current_y)
                current_x = x_start_left
                is_left_column = True

    with open(arquivo_csv, mode="r", encoding="utf-8") as file:
        reader = csv.reader(file)
        next(reader)  # Pula o cabeçalho
        num_questao = 1
        for row in reader:
            # Adiciona o enunciado com quebra de linha automática
            enunciado = f"{num_questao}. {row[5]}"
            enunciado = enunciado.encode("latin-1", "replace").decode("latin-1")
            pdf.set_font("Arial", size=10)  # Fonte menor para ajuste compacto

            # Calcula a altura necessária para o enunciado e ajusta a coluna ou página
            altura_enunciado = (
                pdf.get_string_width(enunciado) // col_width * line_height + line_height
            )
            ajustar_coluna_ou_pagina(altura_enunciado)

            pdf.set_xy(current_x, current_y)
            pdf.multi_cell(col_width, line_height, enunciado)
            current_y = pdf.get_y()  # Atualiza a posição Y após escrever o enunciado

            # Adiciona a imagem da questão, se existir, ajustando a posição
            nome_imagem = row[12].strip()
            imagem_questao = os.path.join(upload_dir, nome_imagem)
            if os.path.exists(imagem_questao):
                tipo_imagem = imagem_questao.split(".")[-1].lower()
                if tipo_imagem in ["jpg", "jpeg", "png", "gif"]:
                    altura_imagem = 45  # Altura estimada para a imagem
                    ajustar_coluna_ou_pagina(altura_imagem)
                    pdf.image(
                        imagem_questao,
                        x=current_x + (col_width - 40) / 2,
                        y=current_y,
                        w=40,
                    )
                    current_y += altura_imagem

            pdf.set_xy(current_x, current_y)

            # Adiciona as alternativas com quebra de linha automática e formatação padrão ENEM
            alternativas = [row[6], row[7], row[8], row[9], row[10]]
            letras = ["A)", "B)", "C)", "D)", "E)"]
            for i, alternativa in enumerate(alternativas):
                if alternativa:
                    alternativa = alternativa.encode("latin-1", "replace").decode(
                        "latin-1"
                    )

                    # Calcula a altura necessária para a alternativa e ajusta a coluna ou página
                    altura_alternativa = (
                        pdf.get_string_width(alternativa)
                        // (col_width - 15)
                        * line_height
                        + line_height
                    )
                    ajustar_coluna_ou_pagina(altura_alternativa)

                    # Define a posição inicial de cada alternativa na mesma coluna
                    pdf.set_xy(current_x, current_y)

                    # Adiciona a letra da alternativa
                    pdf.set_font("Arial", style="B", size=10)
                    pdf.cell(10, line_height, letras[i], align="L")

                    # Adiciona o texto da alternativa alinhado à esquerda
                    pdf.set_font("Arial", size=10)
                    pdf.multi_cell(col_width - 15, line_height, alternativa)
                    current_y = (
                        pdf.get_y()
                    )  # Atualiza a posição Y após cada alternativa

            num_questao += 1

    pdf_file = os.path.join(PDF_FOLDER, f"{nome_avaliacao}.pdf")
    pdf.output(pdf_file)
    return pdf_file


def gerar_docx(arquivo_csv, nome_avaliacao):
    """Gera um arquivo DOCX a partir do arquivo CSV de avaliação com formatação A4."""
    doc = Document()

    # Define as margens da página (em polegadas)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Adiciona imagem de cabeçalho redimensionada para o tamanho A4
    cabecalho_img = os.path.join(IMAGES_DIR, f"{nome_avaliacao}.png")
    if os.path.exists(cabecalho_img):
        # Ajusta a largura para caber na largura de uma página A4 (8.27 inches)
        doc.add_picture(cabecalho_img, width=Inches(4.8))

    # Adiciona o nome da avaliação centralizado
    heading = doc.add_heading(nome_avaliacao, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adiciona questões e formatação
    with open(arquivo_csv, mode="r", encoding="utf-8") as file:
        reader = csv.reader(file)
        num_questao = 1
        for row in reader:
            if reader.line_num > 1:  # Ignora a primeira linha (cabeçalhos)
                # Adiciona enunciado da questão
                enunciado = doc.add_paragraph()
                run = enunciado.add_run(f"{num_questao}. {row[5]}")
                run.font.size = Pt(11)  # Tamanho da fonte
                run.font.name = "Arial"  # Fonte Arial
                run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
                enunciado.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                # Adiciona imagem da questão se existir
                questao_img_path = os.path.join(UPLOAD_FOLDER, row[10])
                if os.path.exists(questao_img_path):
                    doc.add_picture(questao_img_path, width=Inches(4.0))

                # Adiciona alternativas
                alternativas = ["A", "B", "C", "D", "E"]
                for i in range(
                    6, 11
                ):  # Considerando que as alternativas estão nas colunas 6 a 10
                    if row[i]:
                        alternativa_paragrafo = doc.add_paragraph()
                        run = alternativa_paragrafo.add_run(
                            f"({alternativas[i-6]}) {row[i]}"
                        )
                        run.font.size = Pt(11)  # Tamanho da fonte
                        run.font.name = "Arial"  # Fonte Arial
                        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
                        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
                        alternativa_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                num_questao += 1

    # Salva o documento
    docx_file = os.path.join(DOCX_FOLDER, f"{nome_avaliacao}.docx")
    doc.save(docx_file)
    return docx_file
